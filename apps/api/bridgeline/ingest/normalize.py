"""Document type detection, rendering, orientation correction, and deskew."""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from zipfile import BadZipFile, ZipFile

import cv2
import numpy as np
import pypdfium2 as pdfium
import pytesseract
from docx import Document
from PIL import Image, ImageSequence
from pytesseract import Output, TesseractError


class NormalizationError(ValueError):
    """Base class for safe document-normalization failures."""


class OversizeUploadError(NormalizationError):
    """Raised before parsing when an upload exceeds the configured byte limit."""


class UnsupportedDocumentError(NormalizationError):
    """Raised when content sniffing cannot select a safe parser."""


class EncryptedPDFError(NormalizationError):
    """Raised when a password-protected PDF cannot be opened."""


class EmptyDocumentError(NormalizationError):
    """Raised when a supported container has no readable logical pages."""


class MalformedDocumentError(NormalizationError):
    """Raised when a recognized file type cannot be parsed."""


@dataclass(frozen=True, slots=True)
class NormalizedPage:
    """One one-based logical page with visual and embedded-text evidence."""

    number: int
    image_png: bytes | None
    embedded_text: str | None


@dataclass(frozen=True, slots=True)
class NormalizedDocument:
    """Typed output shared by the normalize and OCR stages."""

    filename: str
    media_type: str
    pages: tuple[NormalizedPage, ...]


def normalize_document(
    data: bytes,
    *,
    filename: str,
    max_upload_bytes: int = 50 * 1024 * 1024,
    pdf_dpi: int = 200,
) -> NormalizedDocument:
    """Detect content by signature and normalize it into ordered logical pages."""

    if len(data) > max_upload_bytes:
        raise OversizeUploadError(
            f"Upload exceeds the configured limit of {max_upload_bytes} bytes"
        )
    if not data:
        raise EmptyDocumentError("Document is empty")

    if data.startswith(b"%PDF-"):
        return _normalize_pdf(data, filename=filename, dpi=pdf_dpi)
    if data.startswith(b"PK\x03\x04") and _is_docx(data):
        return _normalize_docx(data, filename=filename)
    if _looks_like_image(data):
        return _normalize_image(data, filename=filename)
    raise UnsupportedDocumentError(
        "Unsupported document type; upload a PDF, DOCX, PNG, JPEG, or TIFF"
    )


def _normalize_pdf(data: bytes, *, filename: str, dpi: int) -> NormalizedDocument:
    try:
        pdf = pdfium.PdfDocument(data)
    except pdfium.PdfiumError as exc:
        message = str(exc).lower()
        if "password" in message or "encrypted" in message:
            raise EncryptedPDFError(
                "This PDF is password-protected; remove encryption and upload it again"
            ) from exc
        raise MalformedDocumentError("The PDF could not be opened") from exc

    try:
        if len(pdf) == 0:
            raise EmptyDocumentError("PDF contains zero pages")
        pages: list[NormalizedPage] = []
        scale = dpi / 72
        for index in range(len(pdf)):
            page = pdf[index]
            embedded = page.get_textpage().get_text_bounded().strip() or None
            image = page.render(scale=scale).to_pil()
            pages.append(
                NormalizedPage(
                    number=index + 1,
                    image_png=_encode_png(prepare_page_image(image)),
                    embedded_text=embedded,
                )
            )
    finally:
        pdf.close()
    return NormalizedDocument(filename=filename, media_type="application/pdf", pages=tuple(pages))


def _normalize_docx(data: bytes, *, filename: str) -> NormalizedDocument:
    try:
        document = Document(BytesIO(data))
    except (ValueError, KeyError, BadZipFile) as exc:
        raise MalformedDocumentError("The DOCX file could not be opened") from exc

    blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    text = "\n".join(blocks).strip()
    if not text:
        raise EmptyDocumentError("DOCX contains no readable pages")
    page = NormalizedPage(number=1, image_png=None, embedded_text=text)
    return NormalizedDocument(
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        pages=(page,),
    )


def _normalize_image(data: bytes, *, filename: str) -> NormalizedDocument:
    try:
        source = Image.open(BytesIO(data))
        media_type = Image.MIME.get(source.format or "", "image/png")
        pages = tuple(
            NormalizedPage(
                number=index,
                image_png=_encode_png(prepare_page_image(frame.copy())),
                embedded_text=None,
            )
            for index, frame in enumerate(ImageSequence.Iterator(source), start=1)
        )
    except (OSError, ValueError) as exc:
        raise MalformedDocumentError("The image could not be opened") from exc
    if not pages:
        raise EmptyDocumentError("Image contains no readable pages")
    return NormalizedDocument(filename=filename, media_type=media_type, pages=pages)


def prepare_page_image(image: Image.Image) -> Image.Image:
    """Deskew small angles and apply Tesseract's coarse orientation result."""

    rgb = image.convert("RGB")
    array = np.asarray(rgb)
    gray = cv2.cvtColor(array, cv2.COLOR_RGB2GRAY)
    threshold = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV | cv2.THRESH_OTSU)[1]
    coordinates = np.column_stack(np.where(threshold > 0))
    if len(coordinates) >= 20:
        angle = cv2.minAreaRect(coordinates[:, ::-1].astype(np.float32))[-1]
        angle = -(90 + angle) if angle < -45 else -angle
        if 0.2 < abs(angle) < 15:
            height, width = array.shape[:2]
            matrix = cv2.getRotationMatrix2D((width / 2, height / 2), angle, 1.0)
            array = cv2.warpAffine(
                array,
                matrix,
                (width, height),
                flags=cv2.INTER_CUBIC,
                borderMode=cv2.BORDER_CONSTANT,
                borderValue=(255, 255, 255),
            )
            rgb = Image.fromarray(array)

    try:
        orientation = pytesseract.image_to_osd(rgb, output_type=Output.DICT)
        rotation = int(orientation.get("rotate", 0)) % 360
    except (TesseractError, ValueError, TypeError):
        rotation = 0
    if rotation:
        rgb = rgb.rotate(-rotation, expand=True, fillcolor="white")
    return rgb


def _encode_png(image: Image.Image) -> bytes:
    stream = BytesIO()
    image.save(stream, format="PNG", optimize=True)
    return stream.getvalue()


def _looks_like_image(data: bytes) -> bool:
    try:
        image = Image.open(BytesIO(data))
        image.verify()
    except (OSError, ValueError):
        return False
    return True


def _is_docx(data: bytes) -> bool:
    try:
        with ZipFile(BytesIO(data)) as archive:
            return "word/document.xml" in archive.namelist()
    except BadZipFile:
        return False
